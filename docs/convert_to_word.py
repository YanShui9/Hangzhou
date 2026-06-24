#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 文档工具
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_word_from_markdown(md_file, docx_file):
    """将Markdown文件转换为Word文档"""

    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    # 添加标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

    # 分割内容为行
    lines = content.split('\n')

    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                # 添加灰色背景（通过段落底纹）
                code_content = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # 处理表格
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_rows = []
            # 跳过分隔行
            if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    table_rows.append(cells)
            i += 1
            # 检查下一行是否还是表格
            if i >= len(lines) or '|' not in lines[i]:
                # 结束表格，创建表格
                if table_rows and len(table_rows) > 0:
                    num_cols = max(len(row) for row in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text
                                # 设置表头样式
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                    doc.add_paragraph()  # 表格后添加空行
                in_table = False
                table_rows = []
            continue

        # 处理标题
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                title_text = match.group(2)
                if level <= 3:
                    doc.add_heading(title_text, level=level)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(title_text)
                    run.bold = True
                    run.font.size = Pt(12)
            i += 1
            continue

        # 处理引用
        if line.startswith('>'):
            quote_text = line[1:].strip()
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # 处理列表
        if re.match(r'^\s*[-*]\s+', line):
            # 无序列表
            list_text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(list_text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\s*\d+\.\s+', line):
            # 有序列表
            list_text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue

        # 处理分割线
        if line.strip() in ['---', '***', '___']:
            doc.add_page_break()
            i += 1
            continue

        # 处理空行
        if line.strip() == '':
            i += 1
            continue

        # 处理普通文本（带格式）
        p = doc.add_paragraph()

        # 处理加粗和斜体
        text = line
        # 加粗
        text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), text)
        # 斜体
        text = re.sub(r'\*(.+?)\*', lambda m: m.group(1), text)
        # 行内代码
        text = re.sub(r'`(.+?)`', lambda m: m.group(1), text)

        # 添加文本，处理格式
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
            else:
                p.add_run(part)

        i += 1

    # 保存文档
    doc.save(docx_file)
    print(f"Word文档已生成：{docx_file}")

if __name__ == '__main__':
    md_file = r'E:\Hangzhou\park-platform\docs\GitHub团队协作指南.md'
    docx_file = r'E:\Hangzhou\park-platform\docs\GitHub团队协作指南.docx'

    create_word_from_markdown(md_file, docx_file)
